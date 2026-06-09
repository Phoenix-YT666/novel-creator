"""
小说导出模块 (Exporter)
将 Markdown 章节导出为 EPUB / PDF / DOCX / TXT / HTML 格式。
"""

from pathlib import Path
from typing import List, Dict, Optional
import json


class Exporter:
    """小说格式导出器"""

    def __init__(self):
        pass

    def to_epub(self, chapters: List[str], metadata: Dict,
               cover: Optional[str] = None, output_path: str = "output.epub") -> str:
        """导出为 EPUB 格式（电子书标准格式）

        Args:
            chapters: 章节目录列表（每个是 Markdown 文本）
            metadata: {title, author, language, description, publisher}
            cover: 封面图片路径
            output_path: 输出路径

        Returns:
            输出文件路径
        """
        try:
            from ebooklib import epub

            book = epub.EpubBook()

            # 元数据
            title = metadata.get("title", "未命名作品")
            author = metadata.get("author", "AI Writer")
            language = metadata.get("language", "zh")

            book.set_identifier(f"novel-{title}-{hash(title)}")
            book.set_title(title)
            book.set_language(language)
            book.add_author(author)

            if metadata.get("description"):
                book.add_metadata("DC", "description", metadata["description"])

            # 封面
            if cover and Path(cover).exists():
                with open(cover, "rb") as f:
                    book.set_cover("cover.jpg", f.read())

            # 样式
            style = """
            body { font-family: serif; line-height: 1.8; }
            h1 { text-align: center; font-size: 1.8em; margin: 2em 0 1em; }
            h2 { font-size: 1.4em; margin: 1.5em 0 0.8em; }
            p { text-indent: 2em; margin: 0.5em 0; }
            """
            css = epub.EpubItem(uid="style", file_name="style/default.css",
                              media_type="text/css", content=style.encode())
            book.add_item(css)

            # 章节
            epub_chapters = []
            spine = ["nav"]

            for i, chapter_text in enumerate(chapters, 1):
                # 提取章节标题
                lines = chapter_text.strip().split("\n")
                chapter_title = lines[0].lstrip("# ") if lines else f"第{i}章"

                # 将 Markdown 转为简单 HTML
                html_content = self._md_to_html(chapter_text)

                epub_chapter = epub.EpubHtml(
                    title=chapter_title,
                    file_name=f"chapter_{i:03d}.xhtml",
                    lang=language,
                )
                epub_chapter.content = f"""<html>
<head><link rel="stylesheet" type="text/css" href="style/default.css"/></head>
<body>{html_content}</body>
</html>"""

                epub_chapter.add_item(css)
                book.add_item(epub_chapter)
                epub_chapters.append(epub_chapter)
                spine.append(epub_chapter)

            # 目录
            book.toc = epub_chapters
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = spine

            # 写入文件
            output = Path(output_path)
            epub.write_epub(str(output), book)

            return str(output.absolute())

        except ImportError:
            print("  ⚠️ ebooklib 未安装。请运行: pip install ebooklib")
            # 降级：写纯文本
            return self.to_txt(chapters, metadata, output_path)

    def to_pdf(self, chapters: List[str], metadata: Dict,
              output_path: str = "output.pdf") -> str:
        """导出为 PDF 格式"""
        html_content = self.to_html(chapters, metadata)

        try:
            from weasyprint import HTML
            HTML(string=html_content).write_pdf(output_path)
        except ImportError:
            print("  ⚠️ weasyprint 未安装。请运行: pip install weasyprint")
            # 降级：保存为 HTML
            output_path = output_path.replace(".pdf", ".html")
            Path(output_path).write_text(html_content, encoding="utf-8")

        return output_path

    def to_docx(self, chapters: List[str], metadata: Dict,
               output_path: str = "output.docx") -> str:
        """导出为 Word 文档"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 标题页
            title = doc.add_heading(metadata.get("title", "未命名作品"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para = doc.add_paragraph(f"作者: {metadata.get('author', 'AI Writer')}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_page_break()

            # 各章节
            for i, chapter_text in enumerate(chapters):
                lines = chapter_text.strip().split("\n")
                chapter_title = lines[0].lstrip("# ").strip() if lines else f"第{i+1}章"

                doc.add_heading(chapter_title, level=1)

                # 正文
                body_text = "\n".join(lines[1:]) if len(lines) > 1 else ""
                for para in body_text.split("\n\n"):
                    para = para.strip()
                    if para:
                        p = doc.add_paragraph(para)
                        p.style.font.size = Pt(12)

                if i < len(chapters) - 1:
                    doc.add_page_break()

            doc.save(output_path)
            return str(Path(output_path).absolute())

        except ImportError:
            print("  ⚠️ python-docx 未安装。请运行: pip install python-docx")
            return self.to_txt(chapters, metadata, output_path)

    def to_txt(self, chapters: List[str], metadata: Dict,
              output_path: str = "output.txt") -> str:
        """导出为纯文本"""
        title = metadata.get("title", "未命名作品")
        author = metadata.get("author", "AI Writer")

        lines = [f"《{title}》", f"作者: {author}", "=" * 50, ""]
        for chapter_text in chapters:
            lines.append(chapter_text)
            lines.append("")
            lines.append("-" * 30)
            lines.append("")

        content = "\n".join(lines)
        Path(output_path).write_text(content, encoding="utf-8")
        return str(Path(output_path).absolute())

    def to_html(self, chapters: List[str], metadata: Dict,
               output_path: str = "output.html") -> str:
        """导出为 HTML"""
        title = metadata.get("title", "未命名作品")
        author = metadata.get("author", "AI Writer")

        html_parts = [f"""<!DOCTYPE html>
<html lang="zh">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>《{title}》</title>
<style>
body {{ max-width: 800px; margin: 0 auto; padding: 20px;
       font-family: 'Noto Serif CJK SC', 'Source Han Serif SC', serif;
       line-height: 2; color: #333; }}
h1 {{ text-align: center; margin: 2em 0 1em; }}
h2 {{ margin: 1.5em 0 0.5em; }}
.toc {{ margin: 2em 0; padding: 1em; background: #f5f5f5; border-radius: 8px; }}
.toc a {{ color: #333; text-decoration: none; }}
p {{ text-indent: 2em; margin: 0.8em 0; }}
.meta {{ text-align: center; color: #999; margin-bottom: 2em; }}
hr {{ border: none; border-top: 1px solid #ddd; margin: 2em 0; }}
</style>
</head>
<body>
<h1>《{title}》</h1>
<p class="meta">{author}</p>

<div class="toc"><h2>📑 目录</h2><ol>
"""]

        # 目录
        for i, chapter_text in enumerate(chapters):
            lines = chapter_text.strip().split("\n")
            chapter_title = lines[0].lstrip("# ").strip() if lines else f"第{i+1}章"
            html_parts.append(f'<li><a href="#ch{i+1}">{chapter_title}</a></li>')

        html_parts.append("</ol></div>")

        # 正文
        for i, chapter_text in enumerate(chapters):
            html_body = self._md_to_html(chapter_text)
            html_parts.append(f'<div id="ch{i+1}">{html_body}</div>')
            if i < len(chapters) - 1:
                html_parts.append("<hr>")

        html_parts.append("</body></html>")

        content = "\n".join(html_parts)

        if output_path:
            Path(output_path).write_text(content, encoding="utf-8")

        return content

    def _md_to_html(self, text: str) -> str:
        """简单的 Markdown → HTML 转换"""
        lines = text.strip().split("\n")
        html_lines = []
        in_paragraph = False

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if in_paragraph:
                    html_lines.append("</p>")
                    in_paragraph = False
                continue

            if stripped.startswith("# "):
                if in_paragraph:
                    html_lines.append("</p>")
                    in_paragraph = False
                html_lines.append(f"<h1>{stripped[2:]}</h1>")
            elif stripped.startswith("## "):
                if in_paragraph:
                    html_lines.append("</p>")
                    in_paragraph = False
                html_lines.append(f"<h2>{stripped[3:]}</h2>")
            elif stripped.startswith("### "):
                if in_paragraph:
                    html_lines.append("</p>")
                    in_paragraph = False
                html_lines.append(f"<h3>{stripped[4:]}</h3>")
            elif stripped == "---" or stripped == "***":
                if in_paragraph:
                    html_lines.append("</p>")
                    in_paragraph = False
                html_lines.append("<hr>")
            else:
                if not in_paragraph:
                    html_lines.append("<p>")
                    in_paragraph = True
                # 简单内联格式
                line_html = stripped
                line_html = line_html.replace("**", "<strong>").replace("__", "</strong>")
                line_html = line_html.replace("*", "<em>").replace("_", "</em>")
                html_lines.append(line_html)

        if in_paragraph:
            html_lines.append("</p>")

        return "\n".join(html_lines)
